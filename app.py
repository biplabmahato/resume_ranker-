import streamlit as st
import os
# from pyresparser import ResumeParser
import spacy
from docx import Document
import pdfminer
import time


import os
import re
from PyPDF2 import PdfReader

# Set up the app
st.set_page_config(page_title="AI Resume Parser", page_icon="📄", layout="wide")

# Initialize NLP
nlp = spacy.load("en_core_web_sm")

# App title and description
st.title("📄 AI Resume Ranking Assistant")
st.markdown("""
Upload job description and candidate resumes to automatically rank candidates based on their qualifications.
""")

# Create directories if they don't exist
os.makedirs("uploads/resumes", exist_ok=True)
os.makedirs("uploads/job_descriptions", exist_ok=True)

# Sidebar for file uploads
with st.sidebar:
    st.header("Upload Files")
    
    # Job description upload
    jd_file = st.file_uploader("Upload Job Description (TXT or DOCX)", type=["txt", "docx"])
    
    # Resume upload (multiple)
    resume_files = st.file_uploader("Upload Resumes (PDF or DOCX)", 
                                  type=["pdf", "docx"], 
                                  accept_multiple_files=True)
    
    # Or use sample data
    use_sample = st.checkbox("Use sample data for demo")
    
    if st.button("Process Files"):
        if (jd_file or use_sample) and (resume_files or use_sample):
            with st.spinner("Analyzing candidates..."):
                time.sleep(2)  # Simulate processing
                st.session_state.processed = True
        else:
            st.warning("Please upload both job description and resumes")

# Function to parse job description
def parse_job_description(text):
    doc = nlp(text)
    skills = []
    experience = 0
    
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PRODUCT", "SKILL"]:
            skills.append(ent.text.lower())
    
    for sent in doc.sents:
        if "experience" in sent.text.lower() and "year" in sent.text.lower():
            try:
                experience = float(''.join(c for c in sent.text if c.isdigit() or c == '.'))
            except:
                pass
    
    return {
        "required_skills": list(set(skills)),
        "experience_required": experience
    }

# Function to parse resume
# def parse_resume(file_path):
#     try:
#         data = ResumeParser(file_path).get_extracted_data()
#         return {
#             "name": data.get("name", "Unknown"),
#             "email": data.get("email", ""),
#             "skills": [s.lower() for s in data.get("skills", [])],
#             "experience": data.get("total_experience", 0),
#             "education": data.get("education", []),
#             "file_name": os.path.basename(file_path)
#         }
#     except Exception as e:
#         st.error(f"Error parsing {file_path}: {str(e)}")
#         return None




def pdf_to_text(file_path):
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

def parse_resume(file_path):
    try:
        # Extract text based on file type
        if file_path.lower().endswith('.pdf'):
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() or ""  # Handle None returns
                    
        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])
        else:  # Assume text file
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

        # Initialize default resume data
        resume_data = {
            "name": os.path.basename(file_path).split('.')[0].replace('_', ' '),  # Use filename as fallback
            "email": "",
            "skills": [],
            "experience": 0,
            "education": [],
            "file_name": os.path.basename(file_path)
        }

        # Extract name - specific pattern for your example resume first
        name_match = re.search(r'#\s*(.+)$', content, re.MULTILINE) or \
                    re.search(r'^([A-Z][a-z]+ [A-Z][a-z]+)', content, re.MULTILINE)
        if name_match:
            resume_data["name"] = name_match.group(1).strip()

        # Extract email - simple pattern
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', content)
        if email_match:
            resume_data["email"] = email_match.group(0).lower()

        # Extract skills - specific pattern for your example
        skills_match = re.search(r'Skills:\s*(.+?)(?:\n\n|$)', content, re.IGNORECASE | re.DOTALL)
        if skills_match:
            skills_text = skills_match.group(1).strip()
            # Split by commas and remove empty items
            skills = [s.strip().lower() for s in re.split(r'[,;]', skills_text) if s.strip()]
            resume_data["skills"] = skills

        # Extract experience - specific pattern for your example
        exp_match = re.search(r'Experience:\s*([\d\.]+)', content, re.IGNORECASE)
        if exp_match:
            try:
                resume_data["experience"] = float(exp_match.group(1))
            except ValueError:
                pass

        # Extract education - specific pattern for your example
        edu_match = re.search(r'Education:\s*(.+?)(?:\n\n|$)', content, re.IGNORECASE | re.DOTALL)
        if edu_match:
            education = edu_match.group(1).strip()
            resume_data["education"] = [education]

        return resume_data

    except Exception as e:
        st.error(f"Error parsing {file_path}: {str(e)}")
        return None
    

# Function to calculate score
def calculate_score(resume_data, jd_requirements):
    score = 0
    
    # Skill matching (50 points)
    resume_skills = set(resume_data["skills"])
    required_skills = set(jd_requirements["required_skills"])
    matched_skills = resume_skills.intersection(required_skills)
    skill_score = 50 * (len(matched_skills) / len(required_skills) if required_skills else 0)
    score += skill_score
    
    # Experience matching (30 points)
    exp_score = 0
    if jd_requirements["experience_required"] > 0:
        exp_ratio = resume_data["experience"] / jd_requirements["experience_required"]
        exp_score = 30 * min(exp_ratio, 1.0)
    else:
        exp_score = 15
    score += exp_score
    
    # Education (20 points)
    edu_score = 20 if any("degree" in edu.lower() or "bachelor" in edu.lower() 
                         for edu in resume_data["education"]) else 0
    score += edu_score
    
    return min(100, round(score, 1)), matched_skills


# Main content area
if "processed" in st.session_state:
    st.header("Results")
    
    # Use sample data if selected
    if use_sample:
        jd_text = """Looking for a Python developer with:
        - 3+ years experience with Django/Flask
        - Knowledge of machine learning
        - Bachelor's degree in Computer Science
        - AWS experience preferred"""
        
        sample_resumes = [
            {"name": "John smith", "skills": ["python", "django", "aws"], 
             "experience": 4, "education": ["Bachelor in Computer Science"],"email":"abc@gamil.com"},
            {"name": "Jane Smith", "skills": ["python", "machine learning"], 
             "experience": 2, "education": ["Master in Data Science"],"email":"bgdugc@gamil.com"},
            {"name": "Alex Johnson", "skills": ["java", "spring"], 
             "experience": 5, "education": ["Bachelor in Software Engineering"],"email":"bccueg@gamil.com"}
        ]
      
    else:
        # Save uploaded files
        if jd_file:
            jd_path = os.path.join("uploads/job_descriptions", jd_file.name)
            with open(jd_path, "wb") as f:
                f.write(jd_file.getbuffer())
            
            if jd_file.type == "text/plain":
                with open(jd_path, "r") as f:
                    jd_text = f.read()
            else:  # DOCX
                doc = Document(jd_path)
                jd_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Process resumes
        sample_resumes = []
        for resume_file in resume_files:
            resume_path = os.path.join("uploads/resumes", resume_file.name)
            with open(resume_path, "wb") as f:
                f.write(resume_file.getbuffer())
            
            parsed = parse_resume(resume_path)
            if parsed:
                sample_resumes.append(parsed)
    
    # Analyze job description
    jd_requirements = parse_job_description(jd_text)
    
    # Display job requirements
    with st.expander("Job Requirements"):
        st.write(jd_text)
        st.subheader("Extracted Requirements")
        st.write(f"**Required Skills:** {', '.join(jd_requirements['required_skills'])}")
        st.write(f"**Years of Experience Required:** {jd_requirements['experience_required']}")
    
    # Score and rank candidates
    ranked_candidates = []
    for resume in sample_resumes:
        score, matched_skills = calculate_score(resume, jd_requirements)
        ranked_candidates.append({
            **resume,
            "score": score,
            "matched_skills": matched_skills
        })
    
    # Sort by score
    ranked_candidates.sort(key=lambda x: x["score"], reverse=True)
    
    # Display results in a nice table
    st.subheader("Ranked Candidates")
    
    for i, candidate in enumerate(ranked_candidates, 1):
        with st.container():
            cols = st.columns([1, 4, 2])
            cols[0].metric(f"Rank #{i}", f"{candidate['score']}/100")
            
            with cols[1]:
                st.subheader(candidate["name"])
                # st.caption(f"📧 {candidate['Email']}")
                st.write(f"**Experience:** {candidate['experience']} years")
                st.write(f"**Education:** {', '.join(candidate['education'])}")
                st.write(f"**Matched Skills:** {', '.join(candidate['matched_skills'])}")
            
            with cols[2]:
                st.progress(candidate["score"] / 100)
                st.caption(f"Overall match score")
                
        st.divider()
    
    # Download results button
#     st.download_button(
#         label="Download Results as CSV",
#         data="\n".join([f"{c['name']},{c['email']},{c['score']}" for c in ranked_candidates]),
#         file_name="candidate_rankings.csv",
#         mime="text/csv"
#     )
# else:
#     st.info("Please upload files and click 'Process Files' in the sidebar")


    # Download results button
    if ranked_candidates:
        # Prepare CSV data
        csv_header = "Name,Email,Score,Matched Skills,Experience,Education\n"
        csv_rows = []
        for candidate in ranked_candidates:
            row = [
                candidate['name'],
                candidate['email'],
                str(candidate['score']),
                ';'.join(candidate['matched_skills']),
                str(candidate['experience']),
                ';'.join(candidate['education'])
            ]
            csv_rows.append(','.join(row))
        
        csv_data = csv_header + '\n'.join(csv_rows)
        
        st.download_button(
            label="Download Results as CSV",
            data=csv_data,
            file_name="candidate_rankings.csv",
            mime="text/csv"
        )
    else:
        st.warning("No candidate data available to download")

# Add some styling
# st.markdown("""
# <style>
#     .stProgress > div > div > div > div {
#         background-color: #4CAF50;
#     }
#     .stMetric {
#         text-align: center;
#     }
# </style>
# """), 